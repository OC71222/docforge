"""Word document extractor using python-docx."""

from __future__ import annotations

from pathlib import Path

from docforge.detector import DocumentFormat
from docforge.extractors.base import BaseExtractor, RawExtraction, RawTable, TextBlock
from docforge.registry import register


@register(DocumentFormat.DOCX)
class DocxExtractor(BaseExtractor):
    def extract(self, file_path: Path, **options: object) -> RawExtraction:
        try:
            from docx import Document
        except ImportError:
            raise ImportError(
                "DOCX support requires python-docx. Install with: "
                "pip install 'docforge[docx]'"
            )

        doc = Document(str(file_path))

        blocks: list[TextBlock] = []
        tables: list[RawTable] = []

        # Extract paragraphs
        y = 0.0
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                y += 12.0
                continue

            is_heading = para.style.name.startswith("Heading")
            heading_level = 0
            if is_heading:
                try:
                    heading_level = int(para.style.name.split()[-1])
                except (ValueError, IndexError):
                    heading_level = 1

            is_bold = any(run.bold for run in para.runs if run.bold is not None)
            font_size = 11.0
            for run in para.runs:
                if run.font.size is not None:
                    font_size = run.font.size.pt
                    break

            blocks.append(TextBlock(
                text=text,
                page=0,
                x0=0,
                y0=y,
                x1=500,
                y1=y + font_size + 2,
                font_size=font_size,
                is_bold=is_bold or is_heading,
                is_heading=is_heading,
                heading_level=heading_level,
            ))
            y += font_size + 6

        # Extract tables
        for table in doc.tables:
            rows_data: list[list[str]] = []
            for row in table.rows:
                rows_data.append([cell.text.strip() for cell in row.cells])

            if not rows_data:
                continue

            headers = rows_data[0]
            data_rows = rows_data[1:] if len(rows_data) > 1 else []

            tables.append(RawTable(
                headers=headers,
                rows=data_rows,
                page=0,
            ))

        # Extract metadata
        props = doc.core_properties
        metadata = {
            "title": props.title or None,
            "author": props.author or None,
            "created_date": str(props.created) if props.created else None,
            "modified_date": str(props.modified) if props.modified else None,
            "page_count": 1,
        }

        return RawExtraction(
            text_blocks=blocks,
            tables=tables,
            images=[],
            metadata=metadata,
            page_count=1,
        )
