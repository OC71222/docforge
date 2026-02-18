"""Shared test fixtures — generates test PDFs programmatically."""

from __future__ import annotations

from pathlib import Path

import pytest

FIXTURES_DIR = Path(__file__).parent / "fixtures"


@pytest.fixture(scope="session", autouse=True)
def generate_fixtures() -> None:
    """Generate test fixtures programmatically."""
    FIXTURES_DIR.mkdir(exist_ok=True)
    _make_simple_pdf()
    _make_tables_pdf()
    _make_multicolumn_pdf()
    _make_sample_html()
    _make_sample_eml()
    _make_sample_docx()
    _make_sample_image()


def _make_simple_pdf() -> None:
    """Create a simple PDF with headings and paragraphs."""
    from fpdf import FPDF

    path = FIXTURES_DIR / "simple.pdf"
    if path.exists():
        return

    pdf = FPDF()
    pdf.add_page()

    pdf.set_font("Helvetica", "B", 24)
    pdf.cell(0, 15, "Document Title", new_x="LMARGIN", new_y="NEXT")

    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 12, "Introduction", new_x="LMARGIN", new_y="NEXT")

    pdf.set_font("Helvetica", "", 11)
    pdf.multi_cell(0, 6, (
        "This is the introduction paragraph. It contains some text that describes "
        "the purpose of this document. DocForge should extract this correctly."
    ))
    pdf.ln(4)

    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 12, "Methods", new_x="LMARGIN", new_y="NEXT")

    pdf.set_font("Helvetica", "", 11)
    pdf.multi_cell(0, 6, (
        "This section describes the methods used. We applied several techniques "
        "to analyze the data and produce meaningful results."
    ))
    pdf.ln(4)

    pdf.set_font("Helvetica", "B", 14)
    pdf.cell(0, 10, "Data Collection", new_x="LMARGIN", new_y="NEXT")

    pdf.set_font("Helvetica", "", 11)
    pdf.multi_cell(0, 6, (
        "Data was collected from multiple sources over a period of six months."
    ))

    pdf.output(str(path))


def _make_tables_pdf() -> None:
    """Create a PDF with a bordered table."""
    from fpdf import FPDF

    path = FIXTURES_DIR / "tables.pdf"
    if path.exists():
        return

    pdf = FPDF()
    pdf.add_page()

    pdf.set_font("Helvetica", "B", 18)
    pdf.cell(0, 12, "Report with Tables", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)

    pdf.set_font("Helvetica", "", 11)
    pdf.multi_cell(0, 6, "The following table shows quarterly results:")
    pdf.ln(4)

    # Table header
    pdf.set_font("Helvetica", "B", 11)
    col_w = 45
    headers = ["Quarter", "Revenue", "Expenses", "Profit"]
    for h in headers:
        pdf.cell(col_w, 8, h, border=1)
    pdf.ln()

    # Table rows
    pdf.set_font("Helvetica", "", 11)
    rows = [
        ["Q1", "$100,000", "$80,000", "$20,000"],
        ["Q2", "$120,000", "$85,000", "$35,000"],
        ["Q3", "$110,000", "$90,000", "$20,000"],
        ["Q4", "$150,000", "$95,000", "$55,000"],
    ]
    for row in rows:
        for cell in row:
            pdf.cell(col_w, 8, cell, border=1)
        pdf.ln()

    pdf.output(str(path))


def _make_multicolumn_pdf() -> None:
    """Create a 2-column layout PDF."""
    from fpdf import FPDF

    path = FIXTURES_DIR / "multicolumn.pdf"
    if path.exists():
        return

    pdf = FPDF()
    pdf.add_page()

    pdf.set_font("Helvetica", "B", 20)
    pdf.cell(0, 14, "Two Column Article", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)

    col_width = 85
    margin = 10
    gutter = 10
    y_start = pdf.get_y()

    # Left column
    pdf.set_xy(margin, y_start)
    pdf.set_font("Helvetica", "B", 14)
    pdf.cell(col_width, 10, "Left Section", new_x="LMARGIN", new_y="NEXT")
    pdf.set_x(margin)
    pdf.set_font("Helvetica", "", 10)
    pdf.multi_cell(col_width, 5, (
        "This is the left column content. It contains information about the first "
        "topic. Readers should see this before the right column content."
    ))

    # Right column
    pdf.set_xy(margin + col_width + gutter, y_start)
    pdf.set_font("Helvetica", "B", 14)
    pdf.cell(col_width, 10, "Right Section", new_x="LMARGIN", new_y="NEXT")
    pdf.set_xy(margin + col_width + gutter, pdf.get_y())
    pdf.set_font("Helvetica", "", 10)
    pdf.multi_cell(col_width, 5, (
        "This is the right column content. It discusses a second topic that "
        "complements the left column."
    ))

    pdf.output(str(path))


def _make_sample_html() -> None:
    """Create a sample HTML file with headings, paragraphs, and a table."""
    path = FIXTURES_DIR / "sample.html"
    if path.exists():
        return

    path.write_text("""\
<!DOCTYPE html>
<html>
<head>
    <title>Sample Document</title>
    <meta name="author" content="Test Author">
</head>
<body>
    <h1>Main Heading</h1>
    <p>This is the first paragraph of the document.</p>
    <h2>Section One</h2>
    <p>Content under section one with some details.</p>
    <table>
        <tr><th>Name</th><th>Value</th></tr>
        <tr><td>Alpha</td><td>100</td></tr>
        <tr><td>Beta</td><td>200</td></tr>
    </table>
    <h2>Section Two</h2>
    <p>Content under section two.</p>
</body>
</html>
""", encoding="utf-8")


def _make_sample_eml() -> None:
    """Create a sample .eml file."""
    path = FIXTURES_DIR / "sample.eml"
    if path.exists():
        return

    path.write_text("""\
From: sender@example.com
To: recipient@example.com
Subject: Test Email Subject
Date: Mon, 01 Jan 2024 12:00:00 +0000
MIME-Version: 1.0
Content-Type: text/plain; charset="utf-8"

This is the body of the test email.

It has multiple paragraphs to verify extraction.

Best regards,
The Sender
""", encoding="utf-8")


def _make_sample_docx() -> None:
    """Create a sample .docx file with headings, paragraphs, and a table."""
    path = FIXTURES_DIR / "sample.docx"
    if path.exists():
        return

    try:
        from docx import Document
    except ImportError:
        return  # Skip if python-docx not installed

    doc = Document()
    doc.core_properties.title = "Sample DOCX"
    doc.core_properties.author = "Test Author"

    doc.add_heading("Document Title", level=1)
    doc.add_paragraph("This is the introduction paragraph.")
    doc.add_heading("Methods", level=2)
    doc.add_paragraph("This section describes the methods used.")

    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Score"
    table.cell(1, 0).text = "Alice"
    table.cell(1, 1).text = "95"
    table.cell(2, 0).text = "Bob"
    table.cell(2, 1).text = "87"

    doc.save(str(path))


def _make_sample_image() -> None:
    """Create a sample image with text for OCR testing."""
    path = FIXTURES_DIR / "sample.png"
    if path.exists():
        return

    try:
        from PIL import Image, ImageDraw, ImageFont
    except ImportError:
        return  # Skip if Pillow not installed

    img = Image.new("RGB", (800, 200), color="white")
    draw = ImageDraw.Draw(img)

    try:
        font = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", 36)
    except (OSError, IOError):
        font = ImageFont.load_default()

    draw.text((50, 30), "Hello World", fill="black", font=font)
    draw.text((50, 100), "DocForge Test Image", fill="black", font=font)

    img.save(str(path))
